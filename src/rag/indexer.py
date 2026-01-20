"""
Document indexer for Azure AI Search.

Handles document processing (PDF, DOCX), chunking, and indexing
to Azure AI Search for the RAG pipeline.
"""

import hashlib
import io
import logging
import re
from dataclasses import dataclass
from datetime import datetime
from typing import Optional

from azure.core.credentials import AzureKeyCredential
from azure.search.documents import SearchClient
from azure.search.documents.indexes import SearchIndexClient
from azure.search.documents.indexes.models import (
    HnswAlgorithmConfiguration,
    SearchableField,
    SearchField,
    SearchFieldDataType,
    SearchIndex,
    SimpleField,
    VectorSearch,
    VectorSearchProfile,
)

from src.config import get_settings
from src.sharepoint import SharePointClient
from src.sharepoint.client import SharePointDocument

from .embeddings import EmbeddingsGenerator

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """Represents a chunk of a document for indexing."""
    
    id: str  # Unique chunk ID
    document_id: str  # Parent document ID
    document_name: str
    chunk_index: int
    content: str
    content_vector: list[float]
    metadata: dict
    created_at: datetime


class DocumentIndexer:
    """
    Indexes documents from SharePoint to Azure AI Search.
    
    Handles:
    - Document fetching from SharePoint
    - Text extraction from PDF and DOCX
    - Intelligent text chunking with overlap
    - Embedding generation
    - Azure AI Search index management
    """
    
    # Chunking parameters
    CHUNK_SIZE = 1000  # Characters per chunk
    CHUNK_OVERLAP = 200  # Overlap between chunks
    
    def __init__(self):
        """Initialize the document indexer."""
        self.settings = get_settings()
        self.sharepoint = SharePointClient()
        self.embeddings = EmbeddingsGenerator()
        
        # Azure AI Search clients
        credential = AzureKeyCredential(
            self.settings.azure_search_api_key.get_secret_value()
        )
        self.index_client = SearchIndexClient(
            endpoint=self.settings.azure_search_endpoint,
            credential=credential,
        )
        self.search_client = SearchClient(
            endpoint=self.settings.azure_search_endpoint,
            index_name=self.settings.azure_search_index_name,
            credential=credential,
        )
    
    async def ensure_index_exists(self) -> None:
        """
        Create the Azure AI Search index if it doesn't exist.
        
        The index schema includes:
        - Text content fields (searchable)
        - Vector field for embeddings
        - Metadata fields for filtering
        """
        index_name = self.settings.azure_search_index_name
        
        try:
            self.index_client.get_index(index_name)
            logger.info(f"Index '{index_name}' already exists")
            return
        except Exception:
            logger.info(f"Creating index '{index_name}'")
        
        # Define vector search configuration
        vector_search = VectorSearch(
            algorithms=[
                HnswAlgorithmConfiguration(
                    name="hnsw-config",
                    parameters={
                        "m": 4,
                        "efConstruction": 400,
                        "efSearch": 500,
                        "metric": "cosine",
                    },
                ),
            ],
            profiles=[
                VectorSearchProfile(
                    name="vector-profile",
                    algorithm_configuration_name="hnsw-config",
                ),
            ],
        )
        
        # Define index fields
        fields = [
            SimpleField(
                name="id",
                type=SearchFieldDataType.String,
                key=True,
            ),
            SimpleField(
                name="document_id",
                type=SearchFieldDataType.String,
                filterable=True,
            ),
            SearchableField(
                name="document_name",
                type=SearchFieldDataType.String,
                filterable=True,
                facetable=True,
            ),
            SimpleField(
                name="chunk_index",
                type=SearchFieldDataType.Int32,
                sortable=True,
            ),
            SearchableField(
                name="content",
                type=SearchFieldDataType.String,
            ),
            SearchField(
                name="content_vector",
                type=SearchFieldDataType.Collection(SearchFieldDataType.Single),
                searchable=True,
                vector_search_dimensions=self.embeddings.dimensions,
                vector_search_profile_name="vector-profile",
            ),
            SimpleField(
                name="created_at",
                type=SearchFieldDataType.DateTimeOffset,
                sortable=True,
            ),
            SearchableField(
                name="source_url",
                type=SearchFieldDataType.String,
            ),
        ]
        
        # Create the index
        index = SearchIndex(
            name=index_name,
            fields=fields,
            vector_search=vector_search,
        )
        
        self.index_client.create_index(index)
        logger.info(f"Created index '{index_name}'")
    
    def extract_text_from_pdf(self, content: bytes) -> str:
        """
        Extract text from a PDF document.
        
        Args:
            content: PDF file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(io.BytesIO(content))
            texts = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    texts.append(text)
            
            return "\n\n".join(texts)
            
        except Exception as e:
            logger.error(f"Failed to extract text from PDF: {e}")
            return ""
    
    def extract_text_from_docx(self, content: bytes) -> str:
        """
        Extract text from a Word document.
        
        Args:
            content: DOCX file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(content))
            texts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    texts.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        texts.append(row_text)
            
            return "\n\n".join(texts)
            
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {e}")
            return ""
    
    def extract_text(self, document: SharePointDocument, content: bytes) -> str:
        """
        Extract text from a document based on its type.
        
        Args:
            document: SharePoint document metadata
            content: Document content as bytes
            
        Returns:
            Extracted text
        """
        if document.is_pdf:
            return self.extract_text_from_pdf(content)
        elif document.is_docx:
            return self.extract_text_from_docx(content)
        else:
            # Try to decode as text
            try:
                return content.decode("utf-8")
            except UnicodeDecodeError:
                logger.warning(f"Unable to extract text from {document.name}")
                return ""
    
    def chunk_text(self, text: str) -> list[str]:
        """
        Split text into overlapping chunks.
        
        Uses sentence boundaries when possible to create
        more coherent chunks.
        
        Args:
            text: Text to chunk
            
        Returns:
            List of text chunks
        """
        # Clean up whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        
        if len(text) <= self.CHUNK_SIZE:
            return [text] if text else []
        
        chunks = []
        start = 0
        
        while start < len(text):
            # Get chunk of approximate target size
            end = start + self.CHUNK_SIZE
            
            if end >= len(text):
                # Last chunk
                chunks.append(text[start:].strip())
                break
            
            # Try to find a sentence boundary near the end
            search_start = max(start + self.CHUNK_SIZE - 100, start)
            search_end = min(start + self.CHUNK_SIZE + 100, len(text))
            search_text = text[search_start:search_end]
            
            # Look for sentence endings
            best_break = -1
            for pattern in ['. ', '.\n', '? ', '?\n', '! ', '!\n', '\n\n']:
                pos = search_text.rfind(pattern)
                if pos != -1:
                    actual_pos = search_start + pos + len(pattern)
                    if actual_pos > start + self.CHUNK_SIZE // 2:
                        best_break = actual_pos
                        break
            
            if best_break == -1:
                # No good sentence break, try word boundary
                search_text = text[start:end + 50]
                space_pos = search_text.rfind(' ')
                if space_pos > self.CHUNK_SIZE // 2:
                    best_break = start + space_pos
                else:
                    best_break = end
            
            chunk = text[start:best_break].strip()
            if chunk:
                chunks.append(chunk)
            
            # Next chunk starts with overlap
            start = best_break - self.CHUNK_OVERLAP
            if start < 0:
                start = best_break
        
        return chunks
    
    async def index_document(
        self,
        document: SharePointDocument,
    ) -> int:
        """
        Index a single document to Azure AI Search.
        
        Args:
            document: SharePoint document to index
            
        Returns:
            Number of chunks indexed
        """
        logger.info(f"Indexing document: {document.name}")
        
        # Download document
        content = await self.sharepoint.download_document(document)
        
        # Extract text
        text = self.extract_text(document, content)
        if not text:
            logger.warning(f"No text extracted from {document.name}")
            return 0
        
        # Chunk text
        chunks = self.chunk_text(text)
        if not chunks:
            return 0
        
        logger.debug(f"Created {len(chunks)} chunks for {document.name}")
        
        # Generate embeddings
        embeddings = await self.embeddings.generate_embeddings_batch(chunks)
        
        # Create chunk documents for indexing
        now = datetime.utcnow()
        documents_to_index = []
        
        for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            chunk_id = hashlib.sha256(
                f"{document.id}_{i}_{chunk[:100]}".encode()
            ).hexdigest()[:32]
            
            documents_to_index.append({
                "id": chunk_id,
                "document_id": document.id,
                "document_name": document.name,
                "chunk_index": i,
                "content": chunk,
                "content_vector": embedding,
                "created_at": now.isoformat() + "Z",
                "source_url": document.web_url,
            })
        
        # Upload to Azure AI Search
        result = self.search_client.upload_documents(documents_to_index)
        
        succeeded = sum(1 for r in result if r.succeeded)
        logger.info(f"Indexed {succeeded}/{len(documents_to_index)} chunks for {document.name}")
        
        return succeeded
    
    async def index_all_documents(self) -> dict[str, int]:
        """
        Index all supported documents from SharePoint.
        
        Returns:
            Dictionary mapping document names to chunk counts
        """
        # Ensure index exists
        await self.ensure_index_exists()
        
        # Get all documents from SharePoint
        documents = await self.sharepoint.list_documents()
        
        # Filter to supported types
        supported_docs = [d for d in documents if d.is_supported]
        logger.info(f"Found {len(supported_docs)} supported documents to index")
        
        results = {}
        for doc in supported_docs:
            try:
                chunk_count = await self.index_document(doc)
                results[doc.name] = chunk_count
            except Exception as e:
                logger.error(f"Failed to index {doc.name}: {e}")
                results[doc.name] = 0
        
        return results
    
    async def delete_document_chunks(self, document_id: str) -> int:
        """
        Delete all chunks for a specific document.
        
        Args:
            document_id: SharePoint document ID
            
        Returns:
            Number of chunks deleted
        """
        # Search for all chunks with this document_id
        results = self.search_client.search(
            search_text="*",
            filter=f"document_id eq '{document_id}'",
            select=["id"],
        )
        
        chunk_ids = [r["id"] for r in results]
        
        if chunk_ids:
            self.search_client.delete_documents([{"id": cid} for cid in chunk_ids])
            logger.info(f"Deleted {len(chunk_ids)} chunks for document {document_id}")
        
        return len(chunk_ids)
